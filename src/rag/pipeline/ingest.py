from __future__ import annotations

import io
from pathlib import PurePosixPath

from clients.postgres_client import PostgresClient
from clients.s3_client import S3Client
from raglog import get_logger
from models import RawDocument
from naming import content_hash

log = get_logger("ingest")


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _parse_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def _parse_markdown(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


_PARSERS = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".doc": _parse_docx,
    ".md": _parse_markdown,
    ".markdown": _parse_markdown,
    ".txt": _parse_markdown,
}


def parse_document(key: str, data: bytes) -> str:
    suffix = PurePosixPath(key).suffix.lower()
    parser = _PARSERS.get(suffix)
    if parser is None:
        raise ValueError(f"unsupported document type: {suffix} ({key})")
    return parser(data)


def _doc_id_from_key(tenant_id: str, key: str) -> str:
    rel = key[len(f"{tenant_id}/"):] if key.startswith(f"{tenant_id}/") else key
    return f"{tenant_id}:{rel}"


class Ingestor:
    """Reads raw docs from object storage, parses them, and records metadata in PostgreSQL."""

    def __init__(self, s3: S3Client | None = None, pg: PostgresClient | None = None) -> None:
        self._s3 = s3 or S3Client()
        self._pg = pg or PostgresClient()

    def ingest_tenant(self, tenant_id: str) -> list[RawDocument]:
        docs: list[RawDocument] = []
        for key in self._s3.list_documents(tenant_id):
            try:
                raw = self._s3.get_bytes(key)
                text = parse_document(key, raw)
            except Exception as exc:  # skip unparseable, keep pipeline moving
                log.warning("parse_failed", key=key, error=str(exc))
                continue
            doc_id = _doc_id_from_key(tenant_id, key)
            docs.append(
                RawDocument(
                    doc_id=doc_id,
                    tenant_id=tenant_id,
                    key=key,
                    content=text,
                    content_hash=content_hash(text),
                )
            )
        log.info("tenant_ingested", tenant_id=tenant_id, docs=len(docs))
        return docs
